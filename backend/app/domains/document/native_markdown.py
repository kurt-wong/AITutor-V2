"""
Native PDF L1 生成器 — PyMuPDF 文本层提取 → L1Document。

按阅读顺序提取文本块，生成 L1Line/L1Page/L1Image，
图片 bbox 通过 get_image_rects(xref) 获取。

详见 Docs/01_Product/T3_IMPLEMENTATION.md §8 Task 1.1。
遵守 V1_LESSONS 3.27（图片 bbox 获取方式）。
"""

from __future__ import annotations

import logging
from pathlib import Path

from app.domains.document.l1_postprocessor import postprocess_l1
from app.domains.document.schemas_l1 import L1Document, L1Image, L1Line, L1Page

logger = logging.getLogger(__name__)

# 页眉页脚过滤阈值
_PAGE_HEADER_Y_MAX = 15
_PAGE_FOOTER_Y_MIN = 800


def extract_l1_from_pdf(
    pdf_path: Path,
    *,
    filename: str | None = None,
    page_range: tuple[int, int] | None = None,
) -> L1Document:
    """从 PDF 提取 L1Document。

    Args:
        pdf_path: PDF 文件路径
        filename: 文件名（默认使用 Path.name）
        page_range: 页码范围 (start, end)，1-based，包含两端。None 表示全部页面。

    Returns:
        L1Document：后处理后的 L1 文档
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF 未安装，请运行: pip install pymupdf"
        )

    doc = fitz.open(str(pdf_path))
    total_pdf_pages = doc.page_count
    fname = filename or pdf_path.name

    # 确定提取范围
    if page_range:
        start_idx = max(0, page_range[0] - 1)
        end_idx = min(total_pdf_pages, page_range[1])
    else:
        start_idx = 0
        end_idx = total_pdf_pages

    pages: list[L1Page] = []
    all_lines: list[L1Line] = []
    all_images: list[L1Image] = []
    global_order = 1

    for page_idx in range(start_idx, end_idx):
        page = doc[page_idx]
        page_no = page_idx + 1  # 1-based

        # 提取文本块
        text_dict = page.get_text("dict")
        blocks = text_dict["blocks"]

        # 分离文本块和图片块
        content_blocks: list[dict] = []
        page_image_xrefs: set[int] = set()

        for block in blocks:
            bbox = block["bbox"]

            if block["type"] == 0:  # 文本块
                # 过滤页眉页脚
                if bbox[1] > _PAGE_FOOTER_Y_MIN or bbox[3] < _PAGE_HEADER_Y_MAX:
                    continue

                # 合并 spans 为块文本
                parts: list[str] = []
                for line_info in block.get("lines", []):
                    line_text = "".join(
                        s["text"] for s in line_info.get("spans", [])
                    )
                    parts.append(line_text)
                text = " ".join(parts).strip()

                if text and len(text) > 0:
                    content_blocks.append({
                        "bbox": bbox,
                        "text": text,
                        "y": bbox[1],
                    })

            elif block["type"] == 1:  # 图块
                if bbox[1] > _PAGE_FOOTER_Y_MIN:
                    continue

                # 获取图片 xref
                img_list = page.get_images(full=True)
                for img_info in img_list:
                    xref = img_info[0]
                    if xref in page_image_xrefs:
                        continue
                    rects = page.get_image_rects(xref)
                    for rect in rects:
                        if (
                            abs(rect.x0 - bbox[0]) < 5
                            and abs(rect.y0 - bbox[1]) < 5
                        ):
                            img_id = f"P{page_no}IMG{len(all_images) + 1:03d}"
                            all_images.append(
                                L1Image(
                                    image_id=img_id,
                                    page_no=page_no,
                                    bbox={
                                        "x1": round(rect.x0, 1),
                                        "y1": round(rect.y0, 1),
                                        "x2": round(rect.x1, 1),
                                        "y2": round(rect.y1, 1),
                                    },
                                    xref=xref,
                                    source="native",
                                    placement="unknown",
                                )
                            )
                            page_image_xrefs.add(xref)
                            break

        # 按 y 坐标排序（阅读顺序）
        content_blocks.sort(key=lambda b: b["y"])

        # 生成 L1Line
        page_lines: list[L1Line] = []
        for block in content_blocks:
            line_no = len(page_lines) + 1
            line = L1Line(
                line_id=f"N{page_no}L{line_no:03d}",
                page_no=page_no,
                line_no_in_page=line_no,
                order=global_order,
                text=block["text"],
                block_type="text",
                bbox={
                    "x1": round(block["bbox"][0], 1),
                    "y1": round(block["bbox"][1], 1),
                    "x2": round(block["bbox"][2], 1),
                    "y2": round(block["bbox"][3], 1),
                },
                source="native",
                continuation=False,
            )
            page_lines.append(line)
            all_lines.append(line)
            global_order += 1

        pages.append(
            L1Page(page_no=page_no, lines=page_lines, images=[])
        )

    doc.close()

    # 计算文本层覆盖率
    text_coverage = _calculate_text_coverage(
        pdf_path, start_idx, end_idx
    )

    # 构建原始 L1Document
    raw_doc = L1Document(
        filename=fname,
        pages=pages,
        lines=all_lines,
        images=all_images,
        source="native",
        total_pages=end_idx - start_idx,
        text_coverage=text_coverage,
    )

    logger.info(
        "native_extract filename=%s pages=%d lines=%d images=%d "
        "text_coverage=%.2f",
        fname,
        end_idx - start_idx,
        len(all_lines),
        len(all_images),
        text_coverage,
    )

    # 执行后处理
    return postprocess_l1(raw_doc)


def _build_numbering_map(doc) -> dict[int, dict[int, dict]]:
    """解析 DOCX numbering part → {numId: {ilvl: {fmt, lvl_text, start}}}。

    2026-08-25：Word 自动编号（如选项 "A."）不在 paragraph.text 中，
    python-docx 提取会丢选项标记（语文 docx Q1 A 选项）。此函数从
    numbering.xml 恢复编号格式，供 extract_l1_from_docx 前缀还原。
    """
    from docx.oxml.ns import qn

    result: dict[int, dict[int, dict]] = {}
    try:
        numbering = doc.part.numbering_part.element
    except Exception:
        return result
    # abstractNumId -> {ilvl: {...}}
    abstracts: dict[str, dict[int, dict]] = {}
    for abs_num in numbering.findall(qn("w:abstractNum")):
        abs_id = abs_num.get(qn("w:abstractNumId"))
        if abs_id is None:
            continue
        lvls: dict[int, dict] = {}
        for lvl in abs_num.findall(qn("w:lvl")):
            try:
                ilvl = int(lvl.get(qn("w:ilvl")) or 0)
            except (TypeError, ValueError):
                continue
            fmt_el = lvl.find(qn("w:numFmt"))
            txt_el = lvl.find(qn("w:lvlText"))
            start_el = lvl.find(qn("w:start"))
            lvls[ilvl] = {
                "fmt": fmt_el.get(qn("w:val")) if fmt_el is not None else "decimal",
                "lvl_text": txt_el.get(qn("w:val")) if txt_el is not None else "%1.",
                "start": int(start_el.get(qn("w:val")) or 1) if start_el is not None else 1,
            }
        abstracts[abs_id] = lvls
    for num in numbering.findall(qn("w:num")):
        num_id_str = num.get(qn("w:numId"))
        abs_el = num.find(qn("w:abstractNumId"))
        if num_id_str is None or abs_el is None:
            continue
        try:
            num_id = int(num_id_str)
        except (TypeError, ValueError):
            continue
        abs_id = abs_el.get(qn("w:val"))
        result[num_id] = abstracts.get(abs_id, {})
    return result


def _format_number(fmt: str, n: int) -> str:
    """按 numbering numFmt 生成序号文本（upperLetter/decimal/lowerLetter/roman）。"""
    if fmt == "upperLetter":
        return chr(64 + n) if 1 <= n <= 26 else str(n)
    if fmt == "lowerLetter":
        return chr(96 + n) if 1 <= n <= 26 else str(n)
    if fmt == "upperRoman":
        return _roman(n).upper()
    if fmt == "lowerRoman":
        return _roman(n)
    return str(n)


def _roman(n: int) -> str:
    vals = [(1000, "m"), (900, "cm"), (500, "d"), (400, "cd"), (100, "c"),
            (90, "xc"), (50, "l"), (40, "xl"), (10, "x"), (9, "ix"),
            (5, "v"), (4, "iv"), (1, "i")]
    out = []
    for v, s in vals:
        while n >= v:
            out.append(s)
            n -= v
    return "".join(out)


def _render_numbering_prefix(lvl: dict, seq: int) -> str:
    """lvlText（如 '%1.'、'（%1）'、'%1)'）替换 %1 为序号。"""
    text = lvl.get("lvl_text") or "%1."
    fmt = lvl.get("fmt") or "decimal"
    start = lvl.get("start") or 1
    label = _format_number(fmt, start + seq - 1)
    return text.replace("%1", label).strip()


def extract_l1_from_docx(
    path: Path,
    *,
    filename: str | None = None,
) -> L1Document:
    """从 DOCX 提取 L1Document（2026-08-25，DOCX 全管线支持）。

    DOCX 是原生文本/表格/图片，不需要 OCR：用 python-docx 直接读取
    段落与表格文本（按文档 body 顺序），生成 L1 行（line_id 用 "D"
    前缀，D1Lxxx，与 PDF 的 N/P 前缀区分）。

    图片：docx 内嵌图片无可提取文本，图片行缺失（图片内容如图表/题干
    插图不影响文本识别，效果待样本验证）。

    Args:
        path: DOCX 文件路径
        filename: 文件名（默认使用 Path.name）

    Returns:
        L1Document：后处理后的 L1 文档
    """
    import logging

    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    fname = filename or path.name
    doc = Document(str(path))
    num_map = _build_numbering_map(doc)
    # 同 (numId, ilvl) 段落序号计数器（从 start 起递增）
    num_seq: dict[tuple[int, int], int] = {}

    lines: list[L1Line] = []
    images: list[L1Image] = []
    order = 1
    line_no = 1

    def _number_prefix(paragraph) -> str:
        """段落自动编号前缀（如 "A."、"1."、"（1）"）；无编号返回 ""。"""
        from docx.oxml.ns import qn

        pPr = paragraph._p.pPr
        if pPr is None:
            return ""
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return ""
        num_id_el = numPr.find(qn("w:numId"))
        ilvl_el = numPr.find(qn("w:ilvl"))
        if num_id_el is None:
            return ""
        try:
            num_id = int(num_id_el.get(qn("w:val")) or 0)
            ilvl = int(ilvl_el.get(qn("w:val")) or 0) if ilvl_el is not None else 0
        except (TypeError, ValueError):
            return ""
        if num_id == 0:  # Word 约定 numId=0 = 关闭编号
            return ""
        lvls = num_map.get(num_id) or {}
        lvl = lvls.get(ilvl)
        if not lvl:
            return ""
        key = (num_id, ilvl)
        seq = num_seq.get(key, 0) + 1
        num_seq[key] = seq
        return _render_numbering_prefix(lvl, seq)

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            para = Paragraph(child, doc)
            prefix = _number_prefix(para)
            text = (prefix + para.text).strip()
            if not text:
                continue
            lines.append(L1Line(
                line_id=f"D1L{line_no:03d}",
                page_no=1,
                line_no_in_page=line_no,
                order=order,
                text=text,
                block_type="text",
                bbox=None,
                source="native",
                continuation=False,
            ))
            line_no += 1
            order += 1
        elif tag.endswith("}tbl"):
            table = Table(child, doc)
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                text = " | ".join(c for c in cells if c)
                if not text:
                    continue
                lines.append(L1Line(
                    line_id=f"D1L{line_no:03d}",
                    page_no=1,
                    line_no_in_page=line_no,
                    order=order,
                    text=text,
                    block_type="table",
                    bbox=None,
                    source="native",
                    continuation=False,
                ))
                line_no += 1
                order += 1

    pages = [L1Page(page_no=1, lines=lines, images=images)]
    raw_doc = L1Document(
        filename=fname,
        pages=pages,
        lines=lines,
        images=images,
        source="native",
        total_pages=1,
        text_coverage=1.0 if lines else 0.0,
    )
    logger.info(
        "docx_extract filename=%s lines=%d images=%d",
        fname,
        len(lines),
        len(images),
    )
    return postprocess_l1(raw_doc)


def _calculate_text_coverage(
    pdf_path: Path, start_idx: int, end_idx: int
) -> float:
    """计算文本层覆盖率。

    覆盖率 = 有文本的页数 / 总页数。
    用于判断是否需要 OCR fallback。
    """
    try:
        import fitz

        doc = fitz.open(str(pdf_path))
        total = end_idx - start_idx
        if total == 0:
            doc.close()
            return 0.0

        has_text = 0
        for i in range(start_idx, end_idx):
            text = doc[i].get_text().strip()
            if text and len(text) > 10:
                has_text += 1
        doc.close()
        return round(has_text / total, 4)
    except Exception:
        return 0.0
